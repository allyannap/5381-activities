"""AI reporter for ICE & Demographics dashboard using OpenAI.

This script:
- Loads the joined census + Vera dataset from ``census_vera_joined.csv``.
- Builds a compact summary of key state-level metrics.
- Sends that summary to an OpenAI model to generate a narrative report.
- Saves the report as plain text, markdown, and a formatted Word document.

Run from the ``hw1`` directory after you have generated ``census_vera_joined.csv``::

    python ai_reporter_openai.py
"""

from __future__ import annotations

from pathlib import Path
from typing import List

import os

import pandas as pd
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI


ROOT = Path(__file__).resolve().parent
DATA_PATH = ROOT / "census_vera_joined.csv"
OUT_TXT = ROOT / "ice_report.txt"
OUT_MD = ROOT / "ice_report.md"
OUT_DOCX = ROOT / "ice_report.docx"

# Load environment variables from .env in this folder
load_dotenv(ROOT / ".env")

# OpenAI configuration
OPENAI_MODEL = "gpt-4.1-mini"  # adjust if you prefer a different model
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY)


def load_dataset(path: Path = DATA_PATH) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(
            f"Expected joined dataset at {path}. Run hw1.py / join_census_vera.py first."
        )
    df = pd.read_csv(path)
    # Ensure consistent state name column
    if "state_name" not in df.columns:
        raise RuntimeError("Expected 'state_name' column in joined dataset.")
    return df


def _top_bottom_block(df: pd.DataFrame, metric: str, label: str, n: int = 5) -> str:
    """Return markdown listing top/bottom n states for a given metric."""
    if metric not in df.columns:
        return f"### {label}\n(Metric '{metric}' not found in dataset.)\n\n"

    work = df[["state_name", metric]].dropna().copy()
    work[metric] = pd.to_numeric(work[metric], errors="coerce")
    work = work.dropna(subset=[metric])

    if work.empty:
        return f"### {label}\n(No valid data for '{metric}'.)\n\n"

    top = work.nlargest(n, metric)
    bottom = work.nsmallest(n, metric)
    median = work[metric].median()

    def fmt_rows(rows: pd.DataFrame) -> List[str]:
        out: List[str] = []
        for _, r in rows.iterrows():
            value = r[metric]
            if "pct" in metric:
                out.append(f"- {r['state_name']}: {value:.2f}%")
            else:
                out.append(f"- {r['state_name']}: {value:,.0f}")
        return out

    lines: List[str] = [f"## {label}", "", f"**Metric column:** `{metric}`", ""]
    lines.append("**Top states:**")
    lines.extend(fmt_rows(top))
    lines.append("")
    lines.append("**Bottom states:**")
    lines.extend(fmt_rows(bottom))
    lines.append("")
    if "pct" in metric:
        lines.append(f"**Median across all states:** {median:.2f}%")
    else:
        lines.append(f"**Median across all states:** {median:,.0f}")
    lines.append("\n")
    return "\n".join(lines)


def build_summary_markdown(df: pd.DataFrame) -> str:
    """Create a compact markdown summary for the AI model to read."""
    lines: List[str] = [
        "# ICE & Demographics – State Summary",
        "",
        "This summary aggregates state-level census demographics and ICE detention context.",
        "Focus metrics:",
        "- `pct_foreign_born`: percent of residents who are foreign-born",
        "- `pct_non_citizen`: percent of residents who are non-citizens",
        "- `ice_facility_count`: number of ICE detention facilities in the state",
        "",
    ]

    lines.append(_top_bottom_block(df, "pct_foreign_born", "Foreign-born population share"))
    lines.append(_top_bottom_block(df, "pct_non_citizen", "Non-citizen population share"))

    if "ice_facility_count" in df.columns:
        lines.append(_top_bottom_block(df, "ice_facility_count", "ICE detention facility counts"))

    return "\n".join(lines)


def call_openai(summary_md: str) -> str:
    """Send the summary to OpenAI and return the report text."""
    prompt = f"""You are an investigative data journalist writing about U.S. immigration enforcement.

You are given a *summary* of state-level demographics and ICE detention activity.
The summary uses the following metrics:
- pct_foreign_born: percent of residents who are foreign-born
- pct_non_citizen: percent of residents who are not U.S. citizens
- ice_facility_count: number of ICE detention facilities in the state

DATA SUMMARY (markdown):
-----------------------
{summary_md}

TASK:
-----
Write a clear, non-technical report about national patterns in this data.

Requirements:
- 3–5 short paragraphs of narrative analysis aimed at a policy or advocacy audience.
- Then a bullet list of 3–5 key takeaways or questions that a reader should remember.
- Refer to specific states and metrics where helpful, but do not drown the reader in numbers.
- Emphasize state comparisons and any regional patterns you notice.
"""

    completion = client.chat.completions.create(
        model=OPENAI_MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are an investigative data journalist writing about U.S. immigration enforcement.",
            },
            {"role": "user", "content": prompt},
        ],
    )
    return completion.choices[0].message.content.strip()


def save_text_and_markdown(report_text: str) -> None:
    """Save the report as plain text and markdown with a simple header."""
    body = report_text.strip()
    header = "ICE & Demographics – AI Report\n" + "=" * 40 + "\n\n"
    formatted = header + body + "\n"

    OUT_TXT.write_text(formatted, encoding="utf-8")
    OUT_MD.write_text(formatted, encoding="utf-8")


def save_docx(report_text: str) -> None:
    """Save the report as a nicely formatted Word document."""
    doc = Document()

    # Main title
    doc.add_heading("ICE & Demographics – AI Report", level=1)

    # Add a blank line between title and content
    doc.add_paragraph("")

    for line in report_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line.strip())
    doc.save(OUT_DOCX)


def main() -> None:
    print("📊 Loading joined census + Vera dataset...")
    df = load_dataset()

    print("🧮 Building compact summary for AI...")
    summary_md = build_summary_markdown(df)

    print("🤖 Contacting OpenAI model", OPENAI_MODEL)
    try:
        report_text = call_openai(summary_md)
    except Exception as exc:  # noqa: BLE001
        raise SystemExit(f"Error calling OpenAI: {exc}") from exc

    print("💾 Saving report to:")
    save_text_and_markdown(report_text)
    print(f" - {OUT_TXT.name}")
    print(f" - {OUT_MD.name}")

    try:
        save_docx(report_text)
        print(f" - {OUT_DOCX.name}")
    except Exception as exc:  # noqa: BLE001
        print(f"⚠️ Could not save Word document: {exc}")

    print("\n✅ AI reporting complete. You can open the .txt/.md/.docx files in the hw1 folder.\n")


if __name__ == "__main__":  # pragma: no cover
    main()
